# 📑 Report Download Center
st.markdown("---")
st.header("📑 Download Reports")

from io import BytesIO
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
import docx
import openpyxl

if st.button("📄 Generate PDF Report"):
    buffer = BytesIO()
    c = canvas.Canvas(buffer, pagesize=letter)
    c.setFont("Helvetica", 12)
    c.drawString(100, 750, "Insurance Premium Report")
    c.drawString(100, 720, f"Age: {age}, Sex: {sex}")
    c.drawString(100, 700, f"BMI: {bmi} ({bmi_category})")
    c.drawString(100, 680, f"Smoker: {smoker}, Region: {region}")
    c.drawString(100, 660, f"Predicted Premium: ${prediction:,.0f}")
    c.showPage()
    c.save()
    buffer.seek(0)
    st.download_button("⬇️ Download PDF", data=buffer, file_name="insurance_report.pdf", mime="application/pdf")

if st.button("📊 Export to Excel"):
    buffer = BytesIO()
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Insurance Report"
    ws.append(["Age","Sex","BMI","BMI Category","Children","Smoker","Region","Predicted Premium"])
    ws.append([age, sex, bmi, bmi_category, children, smoker, region, prediction])
    wb.save(buffer)
    buffer.seek(0)
    st.download_button("⬇️ Download Excel", data=buffer, file_name="insurance_report.xlsx", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

if st.button("📝 Export to Word"):
    buffer = BytesIO()
    doc = docx.Document()
    doc.add_heading("Insurance Premium Report", 0)
    doc.add_paragraph(f"Age: {age}")
    doc.add_paragraph(f"Sex: {sex}")
    doc.add_paragraph(f"BMI: {bmi} ({bmi_category})")
    doc.add_paragraph(f"Children: {children}")
    doc.add_paragraph(f"Smoker: {smoker}")
    doc.add_paragraph(f"Region: {region}")
    doc.add_paragraph(f"Predicted Premium: ${prediction:,.0f}")
    doc.save(buffer)
    buffer.seek(0)
    st.download_button("⬇️ Download Word", data=buffer, file_name="insurance_report.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
